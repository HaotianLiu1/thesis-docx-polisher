import argparse
import json
import os
import re
import sys
import time
from datetime import datetime, timezone
from difflib import SequenceMatcher
from email.utils import parsedate_to_datetime
from pathlib import Path

import requests
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

DEFAULT_BASE_URL = os.getenv("OPENAI_BASE_URL", "")
DEFAULT_API_KEY = os.getenv("OPENAI_API_KEY", "")
DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-5.4")
DEFAULT_SKIP_STYLE_KEYS = ["Heading", "标题-图", "标题-表格", "标题-无编号", "参考文献", "TOC"]

SYSTEM_PROMPT = """
你是中文学术论文润色助手。请先判断段落是否存在以下问题：
- AI语气过重、表达机械
- 语句不通顺或衔接生硬
- 连接词过于套路化

如果没有明显问题，请尽量保持原文不变。
如果有问题，只做微调，严格遵守：
1) 不大幅改写，不改变原意，不新增事实。
2) 优先将过渡词和连接词替换为基础、常用词，表达尽量简单直接。
3) 避免机械连接，使用自然衔接；句式可适度变化（简单句/复合句/插入语混合），避免连续整齐短句。
4) 保留术语、缩写、编号、公式、图表编号、引用编号。

你会收到一个 JSON：{"items":[{"id":0,"text":"..."}, ...]}。
请对每个 item 独立判断并返回同样 id 的结果。

只输出 JSON：
{"items":[{"id":0,"need_edit":true/false,"revised_text":"修改后段落"}, ...]}
不要输出任何额外说明。
""".strip()

TOKEN_RE = re.compile(
    r"\s+|[A-Za-z0-9]+(?:[-._/][A-Za-z0-9]+)*|[\u4e00-\u9fff]|[^\w\s]",
    flags=re.UNICODE,
)


class ApiClient:
    def __init__(self, base_url: str, api_key: str, model: str, timeout: int, retries: int, retry_backoff: float):
        self.url = base_url.rstrip("/") + "/v1/chat/completions"
        self.model = model
        self.timeout = timeout
        self.retries = retries
        self.retry_backoff = retry_backoff
        self.http_calls = 0
        self.headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }


def try_parse_json(text: str):
    s = text.strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s)
        s = re.sub(r"\s*```$", "", s)

    try:
        obj = json.loads(s)
        if isinstance(obj, dict):
            return obj
    except Exception:
        pass

    decoder = json.JSONDecoder()
    for i, ch in enumerate(s):
        if ch != "{":
            continue
        try:
            obj, _ = decoder.raw_decode(s[i:])
            if isinstance(obj, dict):
                return obj
        except Exception:
            continue
    return None


def tokenize(text: str):
    return TOKEN_RE.findall(text)


def diff_chunks(old: str, new: str):
    a, b = tokenize(old), tokenize(new)
    sm = SequenceMatcher(a=a, b=b)
    out = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        oa = "".join(a[i1:i2])
        nb = "".join(b[j1:j2])
        if tag == "equal" and nb:
            out.append(("equal", nb))
        elif tag == "delete" and oa:
            out.append(("delete", oa))
        elif tag == "insert" and nb:
            out.append(("insert", nb))
        elif tag == "replace":
            if oa:
                out.append(("delete", oa))
            if nb:
                out.append(("insert", nb))

    merged = []
    for kind, chunk in out:
        if not chunk:
            continue
        if merged and merged[-1][0] == kind:
            merged[-1] = (kind, merged[-1][1] + chunk)
        else:
            merged.append((kind, chunk))
    return merged


def clear_paragraph_content(paragraph):
    pxml = paragraph._p
    for child in list(pxml):
        if child.tag != qn("w:pPr"):
            pxml.remove(child)


def write_diff(paragraph, old: str, new: str):
    chunks = diff_chunks(old, new)
    if not chunks:
        chunks = [("equal", old)]

    clear_paragraph_content(paragraph)
    for kind, text in chunks:
        run = paragraph.add_run(text)
        if kind == "insert":
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        elif kind == "delete":
            run.font.strike = True


def is_target_paragraph(text: str, style: str, min_len: int, skip_style_keys):
    stripped = text.strip()
    if not stripped:
        return False
    if any(key in style for key in skip_style_keys):
        return False
    if len(stripped) < min_len:
        return False
    return True


def heading1_indices(doc: Document, chapter_style_key: str):
    out = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style is not None else ""
        if chapter_style_key in style:
            out.append(i)
    return out


def reference_index(doc: Document, reference_title: str):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == reference_title:
            return i
    return len(doc.paragraphs)


def chapter_range(doc: Document, start_chapter: int, end_chapter: int, chapter_style_key: str, reference_title: str):
    h1 = heading1_indices(doc, chapter_style_key)
    if not h1:
        raise RuntimeError(f"未找到一级标题样式关键字: {chapter_style_key}")
    if start_chapter < 1:
        raise RuntimeError("--start-chapter 必须 >= 1")
    if start_chapter > len(h1):
        raise RuntimeError(
            f"--start-chapter={start_chapter} 超出范围，文档仅检测到 {len(h1)} 个一级标题"
        )

    start = h1[start_chapter - 1]
    ref = reference_index(doc, reference_title)

    if end_chapter <= 0:
        end = ref
    elif end_chapter < start_chapter:
        raise RuntimeError("--end-chapter 不能小于 --start-chapter")
    elif end_chapter < len(h1):
        end = h1[end_chapter]
    else:
        end = ref

    return start, end, h1, ref


def parse_retry_after(header_value: str):
    if not header_value:
        return None
    value = header_value.strip()
    if not value:
        return None
    try:
        return max(0.0, float(value))
    except Exception:
        pass
    try:
        dt = parsedate_to_datetime(value)
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        wait = (dt - datetime.now(timezone.utc)).total_seconds()
        return max(0.0, wait)
    except Exception:
        return None


def build_batch_user_content(batch_items):
    payload = {
        "items": [{"id": item["id"], "text": item["text"]} for item in batch_items]
    }
    return (
        "请按要求润色以下段落批次。输入 JSON 如下：\n"
        f"{json.dumps(payload, ensure_ascii=False)}\n\n"
        "请严格按以下 JSON 返回，不要输出任何额外内容：\n"
        "{\"items\":[{\"id\":0,\"need_edit\":true,\"revised_text\":\"...\"}]}\n"
        "要求：\n"
        "1) 返回中必须覆盖输入中的每个 id，且 id 不变。\n"
        "2) 无需修改时 need_edit=false，并将原文放入 revised_text。\n"
        "3) 只输出 JSON，不要输出 markdown 代码块或解释。"
    )


def call_openai_compatible_batch(session, client: ApiClient, batch_items):
    payload = {
        "model": client.model,
        "temperature": 0.2,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": build_batch_user_content(batch_items)},
        ],
    }

    last_error = None
    for attempt in range(client.retries + 1):
        try:
            client.http_calls += 1
            resp = session.post(client.url, headers=client.headers, json=payload, timeout=client.timeout)
            retryable_http = resp.status_code == 429 or resp.status_code >= 500
            if retryable_http and attempt < client.retries:
                wait = parse_retry_after(resp.headers.get("Retry-After"))
                if wait is None:
                    wait = client.retry_backoff * (2**attempt)
                time.sleep(wait)
                continue

            resp.raise_for_status()
            data = resp.json()
            content = data.get("choices", [{}])[0].get("message", {}).get("content", "")
            return try_parse_json(content)

        except requests.HTTPError:
            raise
        except requests.RequestException as exc:
            last_error = exc
            if attempt < client.retries:
                time.sleep(client.retry_backoff * (2**attempt))
                continue
            raise
        except Exception as exc:
            last_error = exc
            if attempt < client.retries:
                time.sleep(client.retry_backoff * (2**attempt))
                continue
            raise

    if last_error:
        raise last_error
    raise RuntimeError("API 调用失败")


def parse_item_id(raw_id, batch_size: int):
    if isinstance(raw_id, bool):
        return None
    if isinstance(raw_id, int):
        idx = raw_id
    elif isinstance(raw_id, str) and raw_id.strip().lstrip("-").isdigit():
        idx = int(raw_id.strip())
    else:
        return None
    if 0 <= idx < batch_size:
        return idx
    return None


def normalize_batch_results(batch_items, parsed):
    results = {
        item["id"]: {
            "changed": False,
            "revised_text": item["text"],
            "status": "id_missing",
        }
        for item in batch_items
    }

    if not isinstance(parsed, dict):
        for k in results:
            results[k]["status"] = "json_parse_failed"
        return results

    items = parsed.get("items")
    if not isinstance(items, list):
        for k in results:
            results[k]["status"] = "items_missing"
        return results

    seen = set()
    for row in items:
        if not isinstance(row, dict):
            continue

        idx = parse_item_id(row.get("id"), len(batch_items))
        if idx is None or idx in seen:
            continue
        seen.add(idx)

        old = batch_items[idx]["text"]
        revised = str(row.get("revised_text", "")).strip()
        need_edit = bool(row.get("need_edit", False))

        if not revised:
            results[idx] = {
                "changed": False,
                "revised_text": old,
                "status": "empty_revised_text",
            }
            continue

        if (not need_edit) or revised == old:
            results[idx] = {
                "changed": False,
                "revised_text": old,
                "status": "ok",
            }
            continue

        results[idx] = {
            "changed": True,
            "revised_text": revised,
            "status": "ok",
        }

    return results


def print_analysis(doc: Document, chapter_indexes, ref, chapter_style_key):
    print(f"总段落数: {len(doc.paragraphs)}")
    print(f"一级标题数量: {len(chapter_indexes)}")
    print(f"一级标题样式关键字: {chapter_style_key}")
    for n, idx in enumerate(chapter_indexes, 1):
        title = doc.paragraphs[idx].text.strip()
        print(f"第{n}章标题段落[{idx}]: {title}")
    print(f"参考文献段落索引: {ref}")


def write_fail_log(fail_fp, paragraph_index: int, status: str, error: str = ""):
    if not fail_fp:
        return
    item = {
        "paragraph_index": paragraph_index,
        "status": status,
    }
    if error:
        item["error"] = error
    fail_fp.write(json.dumps(item, ensure_ascii=False) + "\n")


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="ignore")

    parser = argparse.ArgumentParser(
        description="通过 OpenAI 兼容 API 对 DOCX 论文分段润色并标注差异（新增红色，删除删除线）"
    )
    parser.add_argument("--input", required=True, help="输入 .docx")
    parser.add_argument("--output", required=True, help="输出 .docx")
    parser.add_argument("--base-url", default=DEFAULT_BASE_URL, help="OpenAI 兼容接口基地址")
    parser.add_argument("--api-key", default=DEFAULT_API_KEY, help="OpenAI 兼容接口 key")
    parser.add_argument("--model", default=DEFAULT_MODEL, help="模型名")

    parser.add_argument("--start-chapter", type=int, default=1, help="开始章（从1开始）")
    parser.add_argument("--end-chapter", type=int, default=5, help="结束章（包含）。设为0表示直到参考文献")
    parser.add_argument("--chapter-style-key", default="Heading 1", help="一级标题样式关键字")
    parser.add_argument("--reference-title", default="参考文献", help="参考文献标题文本")
    parser.add_argument(
        "--skip-style-keys",
        default=",".join(DEFAULT_SKIP_STYLE_KEYS),
        help="需要跳过的样式关键字，逗号分隔",
    )

    parser.add_argument("--paragraphs-per-call", type=int, default=1, help="每次 API 调用处理的段落数")
    parser.add_argument("--min-len", type=int, default=6, help="最小段落长度")
    parser.add_argument("--timeout", type=int, default=180, help="单次API超时秒")
    parser.add_argument("--retries", type=int, default=2, help="API失败重试次数")
    parser.add_argument("--retry-backoff", type=float, default=1.5, help="重试退避基数秒")
    parser.add_argument("--sleep", type=float, default=0.0, help="每批次调用后休眠秒")
    parser.add_argument("--limit", type=int, default=0, help="调试：最多处理多少个候选段落")
    parser.add_argument("--fail-log", default="", help="保存失败段落日志(jsonl)")
    args = parser.parse_args()

    if not args.base_url:
        raise RuntimeError("缺少 --base-url（或环境变量 OPENAI_BASE_URL）")
    if not args.api_key:
        raise RuntimeError("缺少 --api-key（或环境变量 OPENAI_API_KEY）")
    if args.paragraphs_per_call < 1:
        raise RuntimeError("--paragraphs-per-call 必须 >= 1")

    skip_style_keys = [x.strip() for x in args.skip_style_keys.split(",") if x.strip()]

    in_path = Path(args.input)
    out_path = Path(args.output)

    if not in_path.exists():
        raise FileNotFoundError(f"输入文件不存在: {in_path}")

    try:
        doc = Document(str(in_path))
    except Exception as exc:
        raise RuntimeError(f"无法读取输入文件: {in_path} ({type(exc).__name__}: {exc})") from exc

    start, end, chapter_indexes, ref = chapter_range(
        doc,
        args.start_chapter,
        args.end_chapter,
        args.chapter_style_key,
        args.reference_title,
    )
    print_analysis(doc, chapter_indexes, ref, args.chapter_style_key)
    print(f"处理范围: [{start}, {end})")

    candidate_records = []
    skipped = 0
    for i in range(start, end):
        paragraph = doc.paragraphs[i]
        style = paragraph.style.name if paragraph.style is not None else ""
        old = paragraph.text

        if not is_target_paragraph(old, style, args.min_len, skip_style_keys):
            skipped += 1
            continue

        candidate_records.append(
            {
                "paragraph_index": i,
                "paragraph": paragraph,
                "text": old,
            }
        )

    if args.limit > 0:
        candidate_records = candidate_records[: args.limit]

    total = len(candidate_records)
    batch_count = (total + args.paragraphs_per_call - 1) // args.paragraphs_per_call if total else 0
    print(f"候选段落: {total}，每次 {args.paragraphs_per_call} 段，预计批次: {batch_count}")

    client = ApiClient(
        base_url=args.base_url,
        api_key=args.api_key,
        model=args.model,
        timeout=args.timeout,
        retries=args.retries,
        retry_backoff=args.retry_backoff,
    )

    session = requests.Session()
    edited = 0
    api_err = 0
    fallback_err = 0
    processed = 0

    fail_fp = None
    if args.fail_log:
        fail_path = Path(args.fail_log)
        fail_path.parent.mkdir(parents=True, exist_ok=True)
        fail_fp = fail_path.open("w", encoding="utf-8")

    try:
        for offset in range(0, total, args.paragraphs_per_call):
            batch_records = candidate_records[offset : offset + args.paragraphs_per_call]
            batch_items = [
                {"id": local_id, "text": record["text"]}
                for local_id, record in enumerate(batch_records)
            ]

            try:
                parsed = call_openai_compatible_batch(session, client, batch_items)
            except Exception as exc:
                err = f"{type(exc).__name__}: {exc}"
                for record in batch_records:
                    api_err += 1
                    write_fail_log(fail_fp, record["paragraph_index"], "api_error", err)
                print(f"[API失败][batch {offset // args.paragraphs_per_call + 1}] {err}")
                processed += len(batch_records)
                if processed % 20 == 0 or processed == total:
                    print(f"已处理: {processed}/{total}，已修改: {edited}，API失败: {api_err}，回退: {fallback_err}")
                if args.sleep > 0 and processed < total:
                    time.sleep(args.sleep)
                continue

            normalized = normalize_batch_results(batch_items, parsed)
            for local_id, record in enumerate(batch_records):
                result = normalized.get(local_id)
                if not result:
                    result = {
                        "changed": False,
                        "revised_text": record["text"],
                        "status": "id_missing",
                    }

                if result["changed"]:
                    write_diff(record["paragraph"], record["text"], result["revised_text"])
                    edited += 1

                if result["status"] != "ok":
                    fallback_err += 1
                    write_fail_log(fail_fp, record["paragraph_index"], result["status"])

            processed += len(batch_records)
            if processed % 20 == 0 or processed == total:
                print(f"已处理: {processed}/{total}，已修改: {edited}，API失败: {api_err}，回退: {fallback_err}")

            if args.sleep > 0 and processed < total:
                time.sleep(args.sleep)
    finally:
        if fail_fp:
            fail_fp.close()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    print("\n处理完成")
    print(f"候选段落: {total}")
    print(f"实际修改: {edited}")
    print(f"跳过段落: {skipped}")
    print(f"API失败: {api_err}")
    print(f"协议回退: {fallback_err}")
    print(f"HTTP请求次数: {client.http_calls}")
    print(f"输出文件: {out_path}")


if __name__ == "__main__":
    main()
